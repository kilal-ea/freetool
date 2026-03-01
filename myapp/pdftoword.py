# back/myproject/myapp/pdftoword.py
import json
import logging
import os
import tempfile
from datetime import datetime
from io import BytesIO

import pdfplumber
from docx import Document  # python-docx
from django.http import FileResponse, JsonResponse
from django.views.decorators.csrf import csrf_exempt
from rest_framework.decorators import api_view, parser_classes
from rest_framework.parsers import MultiPartParser, FormParser

from .libreoffice import run_convert, find_converted_file

logger = logging.getLogger(__name__)

def _convert_with_libreoffice(pdf_path, output_dir):
    base_name = os.path.splitext(os.path.basename(pdf_path))[0]
    _, error = run_convert(
        pdf_path,
        output_dir,
        "docx",
        infilter="writer_pdf_import",
        timeout=180,
    )
    if error:
        logger.warning(f"LibreOffice conversion failed: {error}")
        return None

    return find_converted_file(output_dir, base_name, ["docx"])


def _convert_with_pdf2docx(pdf_path, docx_path, settings):
    try:
        from pdf2docx import Converter
    except Exception:
        return False

    try:
        converter = Converter(pdf_path)
        # pdf2docx has limited options, keep defaults for stability
        converter.convert(docx_path, start=0, end=None)
        converter.close()
        return os.path.exists(docx_path)
    except Exception:
        return False


def _build_docx_with_pdfplumber(pdf_bytes, settings):
    extract_tables = settings.get("extractTables", True)
    extract_text = settings.get("extractText", True)

    document = Document()

    with pdfplumber.open(BytesIO(pdf_bytes)) as pdf:
        for i, page in enumerate(pdf.pages):
            document.add_heading(f"Page {i + 1}", level=2)

            if extract_text:
                text = page.extract_text()
                if text:
                    document.add_paragraph(text)

            if extract_tables:
                tables = page.extract_tables()
                for table_data in tables:
                    if table_data:
                        num_columns = len(table_data[0]) if table_data and table_data[0] else 0
                        if num_columns == 0:
                            continue

                        valid_table_data = [
                            row for row in table_data
                            if isinstance(row, list) and len(row) == num_columns
                        ]

                        if not valid_table_data:
                            continue

                        table_doc = document.add_table(
                            rows=len(valid_table_data),
                            cols=num_columns,
                        )
                        table_doc.style = "Table Grid"

                        for row_idx, row_cells in enumerate(valid_table_data):
                            for col_idx, cell_text in enumerate(row_cells):
                                table_doc.cell(row_idx, col_idx).text = str(cell_text or "")
                        document.add_paragraph()

    return document


@api_view(["POST"])
@parser_classes([MultiPartParser, FormParser])
def convert_pdf_to_word(request):
    try:
        pdf_file = request.FILES.get('pdf_file')
        if not pdf_file:
            return JsonResponse({'success': False, 'error': 'No PDF file provided.'}, status=400)

        settings = request.POST.get('settings', '{}')
        try:
            settings = json.loads(settings)
        except json.JSONDecodeError:
            settings = {}

        output_format = settings.get("outputFormat", "docx").lower()
        if output_format != "docx":
            output_format = "docx"
        pdf_bytes = pdf_file.read()

        with tempfile.TemporaryDirectory() as tmp_dir:
            pdf_path = os.path.join(tmp_dir, "input.pdf")
            with open(pdf_path, "wb") as f:
                f.write(pdf_bytes)

            docx_path = os.path.join(tmp_dir, "output.docx")
            libreoffice_docx = _convert_with_libreoffice(pdf_path, tmp_dir)
            if libreoffice_docx and os.path.exists(libreoffice_docx):
                with open(libreoffice_docx, "rb") as f:
                    word_buffer = BytesIO(f.read())
                word_buffer.seek(0)
            else:
                converted = _convert_with_pdf2docx(pdf_path, docx_path, settings)
                if converted:
                    with open(docx_path, "rb") as f:
                        word_buffer = BytesIO(f.read())
                    word_buffer.seek(0)
                else:
                    document = _build_docx_with_pdfplumber(pdf_bytes, settings)
                    word_buffer = BytesIO()
                    document.save(word_buffer)
                    word_buffer.seek(0)

        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f'{pdf_file.name.split(".")[0]}_{timestamp}.{output_format}'

        content_type = (
            "application/msword"
            if output_format == "doc"
            else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        response = FileResponse(
            word_buffer,
            content_type=content_type,
            as_attachment=True,
            filename=filename,
        )
        response['Content-Length'] = word_buffer.getbuffer().nbytes
        response['Access-Control-Expose-Headers'] = 'Content-Disposition'
        return response

    except Exception as e:
        logger.error(f"Error converting PDF to Word: {e}")
        return JsonResponse({'success': False, 'error': f'Failed to convert PDF to Word: {e}'}, status=500)
